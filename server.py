#!/usr/bin/env python3
"""
美的项目客服看板 - 后端 API 服务器
支持本地局域网访问和云部署（Render/Railway/Fly.io 等）

v2: SQLite 持久化存储，支持实时登记、跟进、闭环管理
"""
import os
import sys
import json
import re
import uuid
import sqlite3
import tempfile
import threading
from datetime import datetime, timedelta

try:
    from flask import Flask, request, jsonify, send_from_directory, g
except ImportError:
    print("请先安装 Flask: pip3 install flask")
    sys.exit(1)

import openpyxl
import openpyxl.styles

app = Flask(__name__, static_folder=None)

# CORS — allow access from GitHub Pages and any origin
# Also prevent all caching — data must always be fresh
@app.after_request
def add_cors_headers(response):
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, DELETE, OPTIONS, PATCH'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type, X-Operator'
    response.headers['Access-Control-Max-Age'] = '86400'
    # Prevent caching for API and HTML — data and code must always be fresh
    if request.path.startswith('/api/') or request.path == '/' or request.path.endswith('.html'):
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    return response


def get_operator():
    """Extract operator name from JSON body, URL param, or X-Operator header."""
    from urllib.parse import unquote
    # 1) Try JSON body (_operator field)
    try:
        data = request.get_json(silent=True)
        if data and '_operator' in data:
            name = str(data['_operator']).strip()
            if name:
                return name
    except Exception:
        pass
    # 2) Try URL query param (used by DELETE which has no body)
    try:
        qp = request.args.get('_operator', '')
        if qp:
            name = unquote(qp).strip()
            if name:
                return name
    except Exception:
        pass
    # 3) Try X-Operator header (legacy, URI-encoded)
    try:
        header_val = request.headers.get('X-Operator', '')
        if header_val:
            name = unquote(header_val).strip()
            if name:
                return name
    except Exception:
        pass
    return '匿名'

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_FILE = os.path.join(BASE_DIR, 'dashboard.db')
DATA_FILE = os.path.join(BASE_DIR, 'data.json')
PORT = int(os.environ.get('PORT', 8888))

# ==================== DATABASE ====================

def get_db():
    """Get a thread-local database connection."""
    if 'db' not in g:
        g.db = sqlite3.connect(DB_FILE)
        g.db.row_factory = sqlite3.Row
        g.db.execute("PRAGMA journal_mode=WAL")
        g.db.execute("PRAGMA foreign_keys=ON")
    return g.db


@app.teardown_appcontext
def close_db(exception):
    db = g.pop('db', None)
    if db is not None:
        db.close()


def init_db():
    """Create tables if they don't exist."""
    db = sqlite3.connect(DB_FILE)
    db.execute("PRAGMA journal_mode=WAL")
    db.executescript("""
        CREATE TABLE IF NOT EXISTS daily_records (
            id TEXT PRIMARY KEY,
            date TEXT NOT NULL,
            year TEXT NOT NULL,
            month INTEGER NOT NULL,
            day INTEGER NOT NULL,
            totalDeclared INTEGER DEFAULT 0,
            totalReleased INTEGER DEFAULT 0,
            reviewCompleted INTEGER DEFAULT 0,
            unclearedReason TEXT DEFAULT '',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            source TEXT DEFAULT 'excel'
        );

        CREATE TABLE IF NOT EXISTS abnormal_records (
            id TEXT PRIMARY KEY,
            seq INTEGER,
            date TEXT NOT NULL,
            category TEXT DEFAULT '',
            bizUnit TEXT DEFAULT '',
            company TEXT DEFAULT '',
            importExport TEXT DEFAULT '',
            customsNo TEXT DEFAULT '',
            bolNo TEXT DEFAULT '',
            containerNo TEXT DEFAULT '',
            description TEXT DEFAULT '',
            responsible TEXT DEFAULT '',
            fee TEXT DEFAULT '',
            progress TEXT DEFAULT '',
            status TEXT DEFAULT '',
            agent TEXT DEFAULT '',
            follow_up_notes TEXT DEFAULT '[]',
            year TEXT DEFAULT '',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            source TEXT DEFAULT 'excel',
            deleted INTEGER DEFAULT 0
        );

        CREATE TABLE IF NOT EXISTS operations (
            id TEXT PRIMARY KEY,
            timestamp TEXT NOT NULL,
            action TEXT NOT NULL,
            record_type TEXT DEFAULT '',
            record_id TEXT DEFAULT '',
            record_summary TEXT DEFAULT '',
            changes TEXT DEFAULT '{}',
            operator TEXT DEFAULT '匿名'
        );

        CREATE TABLE IF NOT EXISTS mappings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            year TEXT NOT NULL,
            company TEXT NOT NULL,
            bizUnit TEXT DEFAULT '',
            rep TEXT DEFAULT ''
        );

        CREATE INDEX IF NOT EXISTS idx_daily_date ON daily_records(date);
        CREATE INDEX IF NOT EXISTS idx_daily_year ON daily_records(year);
        CREATE INDEX IF NOT EXISTS idx_abnormal_year ON abnormal_records(year);
        CREATE INDEX IF NOT EXISTS idx_abnormal_deleted ON abnormal_records(deleted);
        CREATE INDEX IF NOT EXISTS idx_operations_timestamp ON operations(timestamp);
    """)
    db.commit()
    db.close()


def migrate_from_json():
    """One-time migration from data.json to SQLite."""
    if not os.path.exists(DATA_FILE):
        return

    db = sqlite3.connect(DB_FILE)
    db.execute("PRAGMA journal_mode=WAL")
    db.row_factory = sqlite3.Row

    # Check if data already exists
    existing = db.execute("SELECT COUNT(*) as cnt FROM daily_records").fetchone()
    if existing['cnt'] > 0:
        db.close()
        return

    print("🔄 检测到 data.json，开始迁移到 SQLite...")
    try:
        with open(DATA_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f"⚠️  读取 data.json 失败: {e}")
        db.close()
        return

    now = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
    daily_count = 0
    abnormal_count = 0
    mapping_count = 0

    for year in ['2025', '2026']:
        year_data = data.get(year)
        if not year_data:
            continue

        # Migrate daily records
        for d in year_data.get('dailyData', []):
            rid = str(uuid.uuid4())
            db.execute("""
                INSERT OR IGNORE INTO daily_records
                (id, date, year, month, day, totalDeclared, totalReleased,
                 reviewCompleted, unclearedReason, created_at, updated_at, source)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'excel')
            """, (
                rid,
                d.get('date', ''),
                str(d.get('year', year)),
                d.get('month', 1),
                d.get('day', 1),
                d.get('totalDeclared', 0),
                d.get('totalReleased', 0),
                d.get('reviewCompleted', 0),
                str(d.get('unclearedReason', '')).strip(),
                now, now
            ))
            daily_count += 1

        # Migrate abnormal records
        for i, r in enumerate(year_data.get('abnormalRecords', [])):
            rid = str(uuid.uuid4())
            db.execute("""
                INSERT OR IGNORE INTO abnormal_records
                (id, seq, date, category, bizUnit, company, importExport,
                 customsNo, bolNo, containerNo, description, responsible,
                 fee, progress, status, agent, follow_up_notes, year,
                 created_at, updated_at, source, deleted)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, '[]', ?, ?, ?, 'excel', 0)
            """, (
                rid,
                r.get('seq', i + 1),
                str(r.get('date', '')).strip(),
                str(r.get('category', '')).strip(),
                str(r.get('bizUnit', '')).strip(),
                str(r.get('company', '')).strip(),
                str(r.get('importExport', '')).strip(),
                str(r.get('customsNo', '')).strip(),
                str(r.get('bolNo', '')).strip(),
                str(r.get('containerNo', '')).strip(),
                str(r.get('description', '')).strip(),
                str(r.get('responsible', '')).strip(),
                str(r.get('fee', '')).strip(),
                str(r.get('progress', '')).strip(),
                str(r.get('status', '')).strip(),
                str(r.get('agent', '')).strip(),
                year,
                now, now
            ))
            abnormal_count += 1

        # Migrate mappings
        for m in year_data.get('mapping', []):
            db.execute("""
                INSERT OR IGNORE INTO mappings (year, company, bizUnit, rep)
                VALUES (?, ?, ?, ?)
            """, (
                year,
                str(m.get('company', '')).strip(),
                str(m.get('bizUnit', '')).strip(),
                str(m.get('rep', '')).strip()
            ))
            mapping_count += 1

    # Log migration
    migration_id = str(uuid.uuid4())
    db.execute("""
        INSERT INTO operations (id, timestamp, action, record_type, record_summary, changes, operator)
        VALUES (?, ?, 'migrate', 'system', ?, ?, '系统')
    """, (
        migration_id,
        now,
        f'data.json → SQLite: {daily_count} 天日报, {abnormal_count} 条异常, {mapping_count} 条映射',
        json.dumps({'daily_count': daily_count, 'abnormal_count': abnormal_count, 'mapping_count': mapping_count}, ensure_ascii=False)
    ))

    db.commit()
    db.close()

    # Rename data.json to backup
    backup_path = DATA_FILE + '.bak'
    try:
        os.rename(DATA_FILE, backup_path)
        print(f"✅ 迁移完成！{daily_count} 天日报, {abnormal_count} 条异常 → SQLite")
        print(f"   data.json 已备份为 data.json.bak")
    except OSError:
        print(f"✅ 迁移完成！(备份失败，请手动重命名 data.json)")


# ==================== HELPERS ====================

def log_operation(action, record_type='', record_id='', record_summary='', changes=None, operator='匿名'):
    """Record an operation to the audit log."""
    db = get_db()
    op_id = str(uuid.uuid4())
    now = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
    db.execute("""
        INSERT INTO operations (id, timestamp, action, record_type, record_id, record_summary, changes, operator)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        op_id, now, action, record_type, record_id, record_summary,
        json.dumps(changes or {}, ensure_ascii=False),
        operator
    ))
    db.commit()


def excel_serial_to_date(serial):
    base = datetime(1899, 12, 30)
    return base + timedelta(days=int(serial))


def parse_workbook(filepath):
    """Parse Excel file, return structured data dict by year."""
    wb = openpyxl.load_workbook(filepath, data_only=True)
    sheet_names = wb.sheetnames

    daily_sheet_names = [n for n in sheet_names if '日报关情况' in n]
    abnormal_sheet_names = [n for n in sheet_names if '异常跟踪表' in n]
    mapping_sheet_name = next((n for n in sheet_names if n == 'Sheet3'), None)

    if not daily_sheet_names:
        raise ValueError('未找到"日报关情况"工作表')

    all_years = {}

    # Parse daily sheets
    for sname in daily_sheet_names:
        ws = wb[sname]
        result = []
        for row in ws.iter_rows(min_row=3, max_row=ws.max_row, values_only=True):
            if not row[0] and not row[1]:
                continue
            try:
                serial = int(row[0])
            except (ValueError, TypeError):
                continue
            if serial < 40000:
                continue
            declared = int(row[1] or 0)
            released = int(row[2] or 0)
            if declared == 0 and released == 0:
                continue
            date = excel_serial_to_date(serial)
            result.append({
                'date': date.strftime('%Y-%m-%d'),
                'year': str(date.year),
                'month': date.month,
                'day': date.day,
                'totalDeclared': declared,
                'totalReleased': released,
                'reviewCompleted': int(row[3] or 0),
                'unclearedReason': str(row[4] or '').strip()
            })
        if result:
            year = str(result[0]['year'])
            if year not in all_years:
                all_years[year] = {'dailyData': [], 'abnormalRecords': [], 'mapping': []}
            all_years[year]['dailyData'].extend(result)

    # Parse abnormal sheets
    for sname in abnormal_sheet_names:
        ws = wb[sname]
        result = []
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, values_only=True):
            if not row[0]:
                continue
            # Skip empty placeholder rows: seq exists but no real data
            # Check essential fields (date + company) are non-empty
            date_val = row[1]
            company_val = row[4]
            if (date_val is None or str(date_val).strip() == '') and \
               (company_val is None or str(company_val).strip() == ''):
                continue
            try:
                serial = int(row[1])
                date_str = excel_serial_to_date(serial).strftime('%Y-%m-%d') if serial > 40000 else str(row[1] or '').strip()
            except (ValueError, TypeError):
                date_str = str(row[1] or '').strip()
            result.append({
                'seq': row[0],
                'date': date_str,
                'category': str(row[2] or '').strip(),
                'bizUnit': str(row[3] or '').strip(),
                'company': str(row[4] or '').strip(),
                'importExport': str(row[5] or '').strip(),
                'customsNo': str(row[6] or '').strip(),
                'bolNo': str(row[7] or '').strip(),
                'containerNo': str(row[8] or '').strip(),
                'description': str(row[9] or '').strip(),
                'responsible': str(row[10] or '').strip(),
                'fee': str(row[11] or '').strip(),
                'progress': str(row[12] or '').strip(),
                'status': str(row[13] or '').strip(),
                'agent': str(row[14] or '').strip()
            })
        if result:
            import re
            ym = re.search(r'(\d{4})', sname)
            year = ym.group(1) if ym else '2026'
            if year not in all_years:
                all_years[year] = {'dailyData': [], 'abnormalRecords': [], 'mapping': []}
            all_years[year]['abnormalRecords'].extend(result)

    # Parse mapping
    mapping = []
    if mapping_sheet_name:
        ws = wb[mapping_sheet_name]
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, values_only=True):
            if row[0]:
                mapping.append({
                    'company': str(row[0]).strip(),
                    'bizUnit': str(row[1]).strip(),
                    'rep': str(row[2]).strip()
                })
    for year in all_years:
        all_years[year]['mapping'] = mapping

    return all_years


def row_to_dict(row, table='daily'):
    """Convert sqlite3.Row to dict."""
    if not row:
        return None
    d = dict(row)
    if table == 'abnormal':
        try:
            d['follow_up_notes'] = json.loads(d.get('follow_up_notes', '[]'))
        except (json.JSONDecodeError, TypeError):
            d['follow_up_notes'] = []
    return d


def get_current_data():
    """Return all data in the same format as the old data.json for frontend compatibility."""
    db = get_db()
    result = {}

    for year in ['2025', '2026']:
        daily_rows = db.execute(
            "SELECT * FROM daily_records WHERE year = ? ORDER BY date", (year,)
        ).fetchall()

        abnormal_rows = db.execute(
            "SELECT * FROM abnormal_records WHERE year = ? AND deleted = 0"
            " AND (date != '' AND date IS NOT NULL)"
            " AND (company != '' AND company IS NOT NULL)"
            " ORDER BY seq", (year,)
        ).fetchall()

        mapping_rows = db.execute(
            "SELECT company, bizUnit, rep FROM mappings WHERE year = ?", (year,)
        ).fetchall()

        if daily_rows or abnormal_rows:
            result[year] = {
                'dailyData': [
                    {
                        'id': r['id'],
                        'date': r['date'],
                        'year': r['year'],
                        'month': r['month'],
                        'day': r['day'],
                        'totalDeclared': r['totalDeclared'],
                        'totalReleased': r['totalReleased'],
                        'reviewCompleted': r['reviewCompleted'],
                        'unclearedReason': r['unclearedReason'],
                        'source': r['source']
                    }
                    for r in daily_rows
                ],
                'abnormalRecords': [
                    {
                        'id': r['id'],
                        'seq': r['seq'],
                        'date': r['date'],
                        'category': r['category'],
                        'bizUnit': r['bizUnit'],
                        'company': r['company'],
                        'importExport': r['importExport'],
                        'customsNo': r['customsNo'],
                        'bolNo': r['bolNo'],
                        'containerNo': r['containerNo'],
                        'description': r['description'],
                        'responsible': r['responsible'],
                        'fee': r['fee'],
                        'progress': r['progress'],
                        'status': r['status'],
                        'agent': r['agent'],
                        'follow_up_notes': json.loads(r['follow_up_notes']) if r['follow_up_notes'] else [],
                        'source': r['source'],
                        'created_at': r['created_at'],
                        'updated_at': r['updated_at']
                    }
                    for r in abnormal_rows
                ],
                'mapping': [
                    {'company': m['company'], 'bizUnit': m['bizUnit'], 'rep': m['rep']}
                    for m in mapping_rows
                ]
            }
        else:
            result[year] = None

    return result


# ==================== ROUTES ====================

@app.route('/')
def index():
    return send_from_directory(BASE_DIR, 'index.html')


@app.route('/<path:filename>')
def static_files(filename):
    if filename.startswith('api/'):
        return jsonify({'error': 'Not found'}), 404
    return send_from_directory(BASE_DIR, filename)


@app.route('/api/health', methods=['GET'])
def health():
    """Health check endpoint for monitoring."""
    db = get_db()
    daily_count = db.execute("SELECT COUNT(*) as cnt FROM daily_records").fetchone()['cnt']
    latest = db.execute("SELECT MAX(date) as max_date FROM daily_records").fetchone()['max_date']
    return jsonify({
        'status': 'ok',
        'daily_records': daily_count,
        'latest_date': latest or 'N/A'
    })


@app.route('/api/data', methods=['GET'])
def get_data():
    """Return the current stored data (same format as before for compatibility)."""
    data = get_current_data()
    resp = jsonify(data)
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    resp.headers['Pragma'] = 'no-cache'
    resp.headers['Expires'] = '0'
    return resp


# ---------------- Upload (merge mode) ----------------

@app.route('/api/upload', methods=['POST'])
def upload():
    """Accept Excel file upload, parse, and merge into database.

    Default behavior (merge): New records are added; existing records (by date for daily,
    by customsNo+bizUnit for abnormal) are NOT overwritten if they have manual edits.

    Query param ?mode=replace enables full replacement of the year's data.
    """
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Empty filename'}), 400

    if not file.filename.lower().endswith(('.xlsx', '.xls')):
        return jsonify({'error': 'Invalid file format, expected .xlsx or .xls'}), 400

    mode = request.args.get('mode', 'merge')
    operator = get_operator()

    try:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
        file.save(tmp.name)
        tmp.close()

        parsed = parse_workbook(tmp.name)
        os.unlink(tmp.name)

        db = get_db()
        now = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
        total_daily = 0
        total_abnormal = 0

        for year in ['2025', '2026']:
            year_data = parsed.get(year)
            if not year_data:
                continue

            if mode == 'replace':
                # Full replace: delete existing data for this year
                db.execute("DELETE FROM daily_records WHERE year = ? AND source = 'excel'", (year,))
                db.execute("DELETE FROM abnormal_records WHERE year = ? AND source = 'excel'", (year,))
                db.execute("DELETE FROM mappings WHERE year = ?", (year,))

            # Merge daily records — skip dates that already exist
            for d in year_data.get('dailyData', []):
                existing = db.execute(
                    "SELECT id FROM daily_records WHERE date = ? AND year = ?",
                    (d['date'], year)
                ).fetchone()
                if existing:
                    continue  # Skip existing dates
                rid = str(uuid.uuid4())
                db.execute("""
                    INSERT INTO daily_records
                    (id, date, year, month, day, totalDeclared, totalReleased,
                     reviewCompleted, unclearedReason, created_at, updated_at, source)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'excel')
                """, (
                    rid, d['date'], year, d['month'], d['day'],
                    d['totalDeclared'], d['totalReleased'],
                    d['reviewCompleted'], d['unclearedReason'],
                    now, now
                ))
                total_daily += 1

            # Merge abnormal records — dedup by customsNo + date + company + bolNo
            # For records where customsNo is '未申报', bolNo acts as the real unique key
            for r in year_data.get('abnormalRecords', []):
                customs_no = str(r.get('customsNo', '')).strip()
                rec_date = str(r.get('date', '')).strip()
                rec_company = str(r.get('company', '')).strip()
                bol_no = str(r.get('bolNo', '')).strip()
                if rec_date:
                    existing = db.execute(
                        "SELECT id FROM abnormal_records WHERE customsNo = ? AND date = ? AND company = ? AND bolNo = ? AND deleted = 0",
                        (customs_no, rec_date, rec_company, bol_no)
                    ).fetchone()
                    if existing:
                        continue
                rid = str(uuid.uuid4())
                db.execute("""
                    INSERT INTO abnormal_records
                    (id, seq, date, category, bizUnit, company, importExport,
                     customsNo, bolNo, containerNo, description, responsible,
                     fee, progress, status, agent, follow_up_notes, year,
                     created_at, updated_at, source, deleted)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, '[]', ?, ?, ?, 'excel', 0)
                """, (
                    rid,
                    r.get('seq', 0),
                    rec_date,
                    str(r.get('category', '')).strip(),
                    str(r.get('bizUnit', '')).strip(),
                    str(r.get('company', '')).strip(),
                    str(r.get('importExport', '')).strip(),
                    customs_no,
                    str(r.get('bolNo', '')).strip(),
                    str(r.get('containerNo', '')).strip(),
                    str(r.get('description', '')).strip(),
                    str(r.get('responsible', '')).strip(),
                    str(r.get('fee', '')).strip(),
                    str(r.get('progress', '')).strip(),
                    str(r.get('status', '')).strip(),
                    str(r.get('agent', '')).strip(),
                    year, now, now
                ))
                total_abnormal += 1

            # Merge mappings
            for m in year_data.get('mapping', []):
                company = str(m.get('company', '')).strip()
                if company:
                    existing = db.execute(
                        "SELECT id FROM mappings WHERE company = ? AND year = ?",
                        (company, year)
                    ).fetchone()
                    if not existing:
                        db.execute(
                            "INSERT INTO mappings (year, company, bizUnit, rep) VALUES (?, ?, ?, ?)",
                            (year, company, str(m.get('bizUnit', '')).strip(), str(m.get('rep', '')).strip())
                        )

        db.commit()

        log_operation(
            action='upload',
            record_type='system',
            record_summary=f'{mode}模式: {file.filename} → {total_daily}天日报, {total_abnormal}条异常',
            changes={'mode': mode, 'filename': file.filename, 'daily_added': total_daily, 'abnormal_added': total_abnormal},
            operator=operator
        )

        return jsonify({
            'success': True,
            'years': [y for y in ['2025', '2026'] if parsed.get(y) and parsed[y].get('dailyData')],
            'totalDays': total_daily,
            'totalAbnormal': total_abnormal,
            'message': f'数据上传成功！新增 {total_daily} 天记录，{total_abnormal} 条异常'
        })

    except Exception as e:
        return jsonify({'error': f'文件解析失败: {str(e)}'}), 500


# ---------------- Abnormal Records CRUD ----------------

@app.route('/api/abnormal-records', methods=['POST'])
def create_abnormal_record():
    """Create a new abnormal record manually."""
    data = request.get_json(silent=True)
    if not data:
        return jsonify({'error': 'Invalid JSON'}), 400

    db = get_db()
    now = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
    rid = str(uuid.uuid4())
    rec_date = str(data.get('date', datetime.now().strftime('%Y-%m-%d'))).strip()
    year = rec_date[:4] if len(rec_date) >= 4 else str(datetime.now().year)

    # Auto-assign seq number
    max_seq = db.execute(
        "SELECT COALESCE(MAX(seq), 0) as mx FROM abnormal_records WHERE year = ?", (year,)
    ).fetchone()['mx']
    seq = max_seq + 1

    db.execute("""
        INSERT INTO abnormal_records
        (id, seq, date, category, bizUnit, company, importExport,
         customsNo, bolNo, containerNo, description, responsible,
         fee, progress, status, agent, follow_up_notes, year,
         created_at, updated_at, source, deleted)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, '[]', ?, ?, ?, 'manual', 0)
    """, (
        rid, seq, rec_date,
        str(data.get('category', '')).strip(),
        str(data.get('bizUnit', '')).strip(),
        str(data.get('company', '')).strip(),
        str(data.get('importExport', '')).strip(),
        str(data.get('customsNo', '')).strip(),
        str(data.get('bolNo', '')).strip(),
        str(data.get('containerNo', '')).strip(),
        str(data.get('description', '')).strip(),
        str(data.get('responsible', '')).strip(),
        str(data.get('fee', '')).strip(),
        str(data.get('progress', '')).strip(),
        str(data.get('status', '未闭环')).strip(),
        str(data.get('agent', '')).strip(),
        year, now, now
    ))
    db.commit()

    log_operation(
        action='create', record_type='abnormal', record_id=rid,
        record_summary=f'新增异常: {data.get("company", "")} {data.get("customsNo", "")}',
        changes={'new': {k: v for k, v in data.items() if v}},
        operator=get_operator()
    )

    return jsonify({'success': True, 'id': rid, 'seq': seq}), 201


@app.route('/api/abnormal-records/<record_id>', methods=['PUT'])
def update_abnormal_record(record_id):
    """Update an abnormal record. Only provided fields are updated."""
    data = request.get_json(silent=True)
    if not data:
        return jsonify({'error': 'Invalid JSON'}), 400

    db = get_db()
    now = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')

    old = db.execute("SELECT * FROM abnormal_records WHERE id = ? AND deleted = 0", (record_id,)).fetchone()
    if not old:
        return jsonify({'error': 'Record not found'}), 404

    # Build update — only set fields that are provided and different
    allowed_fields = [
        'date', 'category', 'bizUnit', 'company', 'importExport',
        'customsNo', 'bolNo', 'containerNo', 'description', 'responsible',
        'fee', 'progress', 'status', 'agent', 'seq'
    ]
    updates = {}
    changes = {}
    for field in allowed_fields:
        if field in data:
            new_val = str(data[field]).strip() if data[field] is not None else ''
            old_val = str(old[field] or '').strip()
            if new_val != old_val:
                updates[field] = new_val
                changes[field] = {'old': old_val, 'new': new_val}

    if not updates:
        return jsonify({'success': True, 'message': 'No changes'})

    set_clauses = ', '.join(f'{k} = ?' for k in updates.keys())
    values = list(updates.values()) + [now, record_id]

    db.execute(
        f"UPDATE abnormal_records SET {set_clauses}, updated_at = ? WHERE id = ?",
        values
    )
    db.commit()

    if changes:
        log_operation(
            action='update', record_type='abnormal', record_id=record_id,
            record_summary=f'更新记录: {old["company"]} {old["customsNo"]}',
            changes=changes,
            operator=get_operator()
        )

    return jsonify({'success': True, 'changes': changes})


@app.route('/api/abnormal-records/<record_id>', methods=['DELETE'])
def delete_abnormal_record(record_id):
    """Soft-delete an abnormal record."""
    db = get_db()
    old = db.execute("SELECT * FROM abnormal_records WHERE id = ? AND deleted = 0", (record_id,)).fetchone()
    if not old:
        return jsonify({'error': 'Record not found'}), 404

    db.execute("UPDATE abnormal_records SET deleted = 1, updated_at = ? WHERE id = ?",
               (datetime.now().strftime('%Y-%m-%dT%H:%M:%S'), record_id))
    db.commit()

    log_operation(
        action='delete', record_type='abnormal', record_id=record_id,
        record_summary=f'删除记录: {old["company"]} {old["customsNo"]}',
        operator=get_operator()
    )

    return jsonify({'success': True})


@app.route('/api/abnormal-records/batch-update', methods=['POST'])
def batch_update_abnormal():
    """Batch update status of multiple abnormal records."""
    data = request.get_json(silent=True)
    if not data or 'ids' not in data or 'status' not in data:
        return jsonify({'error': 'Requires ids (array) and status (string)'}), 400

    ids = data['ids']
    new_status = str(data['status']).strip()
    if not ids or not new_status:
        return jsonify({'error': 'ids and status required'}), 400

    db = get_db()
    now = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
    count = 0
    summaries = []

    for rid in ids:
        old = db.execute("SELECT * FROM abnormal_records WHERE id = ? AND deleted = 0", (rid,)).fetchone()
        if not old:
            continue
        old_status = (old['status'] or '').strip()
        if old_status == new_status:
            continue
        db.execute("UPDATE abnormal_records SET status = ?, updated_at = ? WHERE id = ?",
                   (new_status, now, rid))
        summaries.append(f'{old["company"]} {old["customsNo"]}')
        count += 1

    db.commit()

    if count > 0:
        log_operation(
            action='batch_update', record_type='abnormal',
            record_summary=f'批量更新 {count} 条 → {new_status}',
            changes={'ids': ids, 'new_status': new_status, 'count': count, 'summaries': summaries},
            operator=get_operator()
        )

    return jsonify({'success': True, 'updated': count})


# ---------------- Follow-up Notes ----------------

@app.route('/api/abnormal-records/<record_id>/notes', methods=['POST'])
def add_follow_up_note(record_id):
    """Add a follow-up note to an abnormal record."""
    data = request.get_json(silent=True)
    if not data or 'note' not in data:
        return jsonify({'error': 'note is required'}), 400

    db = get_db()
    old = db.execute("SELECT * FROM abnormal_records WHERE id = ? AND deleted = 0", (record_id,)).fetchone()
    if not old:
        return jsonify({'error': 'Record not found'}), 404

    now = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
    note_entry = {
        'date': now,
        'note': str(data['note']).strip(),
        'author': get_operator()
    }

    try:
        existing_notes = json.loads(old['follow_up_notes'] or '[]')
    except json.JSONDecodeError:
        existing_notes = []

    existing_notes.append(note_entry)
    db.execute(
        "UPDATE abnormal_records SET follow_up_notes = ?, updated_at = ? WHERE id = ?",
        (json.dumps(existing_notes, ensure_ascii=False), now, record_id)
    )
    db.commit()

    log_operation(
        action='add_note', record_type='abnormal', record_id=record_id,
        record_summary=f'添加备注: {old["company"]} {old["customsNo"]}',
        changes={'note': note_entry},
        operator=get_operator()
    )

    return jsonify({'success': True, 'notes': existing_notes})


# ---------------- Daily Records ----------------

@app.route('/api/daily-records', methods=['POST'])
def create_daily_record():
    """Manually create/update a daily business record. Supports accumulation and duplicate detection."""
    data = request.get_json(silent=True)
    if not data or 'date' not in data:
        return jsonify({'error': 'date is required'}), 400

    db = get_db()
    now = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
    rec_date = str(data['date']).strip()
    year = rec_date[:4]

    new_reason = str(data.get('unclearedReason', '')).strip()
    new_declared = data.get('totalDeclared', 0)
    new_released = data.get('totalReleased', 0)
    new_review = data.get('reviewCompleted', 0)

    # Check if date already exists (needed early for duplicate check exclusion)
    existing = db.execute(
        "SELECT * FROM daily_records WHERE date = ? AND year = ?",
        (rec_date, year)
    ).fetchone()

    # === Duplicate detection ===
    # When updating same date, exclude own record's IDs so pre-filled items don't
    # cause false duplicates. The merge function handles dedup within the same record.
    if new_reason and new_reason != '无':
        new_ids = _extract_item_ids(new_reason)
        if new_ids:
            all_rows = db.execute(
                "SELECT id, date, unclearedReason FROM daily_records"
                " WHERE unclearedReason IS NOT NULL AND unclearedReason != '' AND unclearedReason != '无'"
            ).fetchall()
            existing_ids = set()  # IDs from OTHER records only
            for row in all_rows:
                if existing and row['id'] == existing['id']:
                    continue  # Skip own record — merge handles dedup
                existing_ids |= _extract_item_ids(row['unclearedReason'] or '')

            duplicates = new_ids & existing_ids
            if duplicates:
                dup_list = '、'.join(sorted(duplicates))
                return jsonify({
                    'error': f'提单号重复，以下单号已录入：{dup_list}',
                    'duplicates': sorted(duplicates)
                }), 409

    if existing:
        # === Accumulate (叠加) instead of overwrite ===
        total_declared = (existing['totalDeclared'] or 0) + new_declared
        total_released = (existing['totalReleased'] or 0) + new_released
        total_review = (existing['reviewCompleted'] or 0) + new_review

        # Merge unclearedReason: append new items to existing sections
        existing_reason = existing['unclearedReason'] or ''
        merged_reason = _merge_uncleared_reason(existing_reason, new_reason)

        db.execute("""
            UPDATE daily_records SET
                totalDeclared = ?, totalReleased = ?, reviewCompleted = ?,
                unclearedReason = ?, updated_at = ?, source = 'manual'
            WHERE id = ?
        """, (total_declared, total_released, total_review, merged_reason, now, existing['id']))
        db.commit()
        log_operation(
            action='update_daily', record_type='daily', record_id=existing['id'],
            record_summary=f'叠加日报: {rec_date} (+{new_declared}申报/+{new_released}放行/+{new_review}审结)',
            changes=data,
            operator=get_operator()
        )
        return jsonify({'success': True, 'id': existing['id'], 'action': 'updated'})

    # Create new
    rid = str(uuid.uuid4())
    dt = datetime.strptime(rec_date, '%Y-%m-%d') if len(rec_date) == 10 else datetime.now()
    db.execute("""
        INSERT INTO daily_records
        (id, date, year, month, day, totalDeclared, totalReleased,
         reviewCompleted, unclearedReason, created_at, updated_at, source)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'manual')
    """, (
        rid, rec_date, year, dt.month, dt.day,
        new_declared, new_released, new_review,
        new_reason if new_reason else '',
        now, now
    ))
    db.commit()

    log_operation(
        action='create_daily', record_type='daily', record_id=rid,
        record_summary=f'新增日报: {rec_date}',
        changes=data,
        operator=get_operator()
    )

    return jsonify({'success': True, 'id': rid, 'action': 'created'}), 201


@app.route('/api/daily-records/by-date/<rec_date>', methods=['GET'])
def get_daily_record_by_date(rec_date):
    """Return a daily record by date for form pre-fill, or null if not found."""
    db = get_db()
    year = rec_date[:4]
    row = db.execute(
        "SELECT * FROM daily_records WHERE date = ? AND year = ?",
        (rec_date, year)
    ).fetchone()

    if not row:
        return jsonify({'found': False, 'record': None})

    return jsonify({
        'found': True,
        'record': {
            'id': row['id'],
            'date': row['date'],
            'totalDeclared': row['totalDeclared'],
            'totalReleased': row['totalReleased'],
            'reviewCompleted': row['reviewCompleted'],
            'unclearedReason': row['unclearedReason']
        }
    })


# ---------------- Clear ----------------

@app.route('/api/clear', methods=['POST'])
def clear_data():
    """Clear all data after confirmation. Requires exact confirmation string."""
    data = request.get_json(silent=True)
    if not data or data.get('confirmation') != '确认清零所有数据':
        return jsonify({'error': 'Confirmation string does not match. Required: 确认清零所有数据'}), 403

    db = get_db()
    now = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')

    # Count before clearing
    daily_count = db.execute("SELECT COUNT(*) as cnt FROM daily_records").fetchone()['cnt']
    abnormal_count = db.execute("SELECT COUNT(*) as cnt FROM abnormal_records WHERE deleted = 0").fetchone()['cnt']

    # Log the clear operation BEFORE deleting
    log_operation(
        action='clear',
        record_type='system',
        record_summary=f'清零所有数据: {daily_count} 天日报, {abnormal_count} 条异常',
        changes={'daily_count': daily_count, 'abnormal_count': abnormal_count},
        operator=get_operator()
    )

    # Delete all data
    db.execute("DELETE FROM daily_records")
    db.execute("DELETE FROM abnormal_records")
    db.execute("DELETE FROM mappings")
    db.commit()

    return jsonify({
        'success': True,
        'message': f'已清零 {daily_count} 天日报, {abnormal_count} 条异常',
        'daily_cleared': daily_count,
        'abnormal_cleared': abnormal_count
    })


# ---------------- Operations Log ----------------

@app.route('/api/operations', methods=['GET'])
def get_operations():
    """Get operations log. Optional ?record_id= to filter by record."""
    limit = request.args.get('limit', 100, type=int)
    record_id = request.args.get('record_id', '').strip()
    db = get_db()
    if record_id:
        rows = db.execute(
            "SELECT * FROM operations WHERE record_id = ? ORDER BY timestamp DESC LIMIT ?",
            (record_id, limit)
        ).fetchall()
    else:
        rows = db.execute(
            "SELECT * FROM operations ORDER BY timestamp DESC LIMIT ?", (limit,)
        ).fetchall()
    return jsonify([dict(r) for r in rows])


# ---------------- Export Report ----------------

@app.route('/api/export-report', methods=['GET'])
def export_report():
    """Export daily report Excel file from database data.

    Uses the latest Excel template for formatting, populates 2026 daily
    and abnormal records from SQLite, and returns only the two 2026 sheets.
    """
    from io import BytesIO
    from urllib.parse import quote

    db = get_db()

    # Query all 2026 daily records
    daily_rows = db.execute(
        "SELECT * FROM daily_records WHERE year = '2026' ORDER BY date"
    ).fetchall()

    # Query all 2026 abnormal records (not deleted), skip empty placeholder rows
    abnormal_rows = db.execute(
        "SELECT * FROM abnormal_records WHERE year = '2026' AND deleted = 0"
        " AND (date != '' AND date IS NOT NULL)"
        " AND (company != '' AND company IS NOT NULL)"
        " ORDER BY seq"
    ).fetchall()

    # Find the latest Excel template for styles
    import glob
    candidates = glob.glob(os.path.join(BASE_DIR, '华东口岸申报日报关情况*.xlsx'))
    if not candidates:
        return jsonify({'error': '模板文件不存在，请上传华东口岸申报日报关情况.xlsx'}), 500
    template_path = max(candidates, key=os.path.getmtime)

    wb = openpyxl.load_workbook(template_path)

    # ============================================================
    # Process "2026 日报关情况" sheet
    # ============================================================
    ws_daily = wb['2026 日报关情况']

    # Unmerge all merged cells in data area (from row 3 onwards, but keep A1:E1 title)
    for mc in list(ws_daily.merged_cells.ranges):
        if mc.min_row >= 3:
            ws_daily.unmerge_cells(str(mc))

    # Capture styles from the first data row (row 3) before deleting
    daily_styles = {}
    for col in range(1, 6):
        cell = ws_daily.cell(row=3, column=col)
        daily_styles[col] = {
            'font': openpyxl.styles.Font(
                name=cell.font.name,
                size=cell.font.size,
                bold=cell.font.bold,
                color=cell.font.color
            ),
            'alignment': openpyxl.styles.Alignment(
                horizontal=cell.alignment.horizontal or 'center',
                vertical=cell.alignment.vertical or 'center',
                wrap_text=(col == 5)  # wrap text for the reason column
            ),
            'number_format': cell.number_format,
            'border': openpyxl.styles.Border(
                left=cell.border.left, right=cell.border.right,
                top=cell.border.top, bottom=cell.border.bottom
            ),
        }

    # Delete all old data rows (rows 3 to end)
    if ws_daily.max_row >= 3:
        ws_daily.delete_rows(3, ws_daily.max_row - 2)

    # Write fresh data from database
    for i, record in enumerate(daily_rows):
        row_num = 3 + i
        date_str = record['date']
        try:
            dt = datetime.strptime(date_str, '%Y-%m-%d')
            serial = (dt - datetime(1899, 12, 30)).days
        except (ValueError, TypeError):
            serial = date_str

        values = [
            serial,
            record['totalDeclared'] or 0,
            record['totalReleased'] or 0,
            record['reviewCompleted'] or 0,
            (record['unclearedReason'] or '').strip() or '无'
        ]

        for col, val in enumerate(values, 1):
            cell = ws_daily.cell(row=row_num, column=col, value=val)
            style = daily_styles.get(col, {})
            if 'font' in style:
                cell.font = style['font']
            if 'alignment' in style:
                cell.alignment = style['alignment']
            if 'border' in style:
                cell.border = style['border']
            # Column A (日期): use explicit date format so Excel displays it properly
            if col == 1:
                cell.number_format = 'YYYY-MM-DD'
            elif 'number_format' in style:
                cell.number_format = style['number_format']

        # Auto-adjust row height based on unclearedReason content
        reason_text = (record['unclearedReason'] or '').strip()
        lines = reason_text.count('\n') + 1 if reason_text and reason_text != '无' else 1
        row_height = max(15, lines * 15.6)  # ~15.6pt per line at 11pt font
        ws_daily.row_dimensions[row_num].height = row_height

    # ============================================================
    # Process "2026年异常跟踪表" sheet
    # ============================================================
    ws_abnormal = wb['2026年异常跟踪表']

    # Unmerge all data-area merged cells (from row 2 onwards)
    for mc in list(ws_abnormal.merged_cells.ranges):
        if mc.min_row >= 2:
            ws_abnormal.unmerge_cells(str(mc))

    # Capture styles from the first data row (row 2) before deleting
    abnormal_styles = {}
    for col in range(1, 17):
        cell = ws_abnormal.cell(row=2, column=col)
        abnormal_styles[col] = {
            'font': openpyxl.styles.Font(
                name=cell.font.name,
                size=cell.font.size,
                bold=cell.font.bold,
                color=cell.font.color
            ),
            'alignment': openpyxl.styles.Alignment(
                horizontal=cell.alignment.horizontal or 'center',
                vertical=cell.alignment.vertical or 'center',
                wrap_text=(col in (10, 13))  # wrap for 异常描述 and 处理进度
            ),
            'number_format': cell.number_format,
            'border': openpyxl.styles.Border(
                left=cell.border.left, right=cell.border.right,
                top=cell.border.top, bottom=cell.border.bottom
            ),
        }

    # Delete all old data rows (rows 2 to end)
    if ws_abnormal.max_row >= 2:
        ws_abnormal.delete_rows(2, ws_abnormal.max_row - 1)

    # Write fresh data from database
    for i, record in enumerate(abnormal_rows):
        row_num = 2 + i
        date_str = record['date']
        try:
            dt = datetime.strptime(date_str, '%Y-%m-%d')
            serial = (dt - datetime(1899, 12, 30)).days
        except (ValueError, TypeError):
            serial = date_str

        values = [
            record['seq'],                # A: 序号 (keep original, don't re-number)
            serial,                        # B: 发生/发现日期
            record['category'] or '',       # C: 异常类别
            record['bizUnit'] or '',        # D: 事业部
            record['company'] or '',        # E: 经营单位
            record['importExport'] or '',   # F: 进出口
            record['customsNo'] or '',      # G: 报关单号
            record['bolNo'] or '',          # H: 提运单号
            record['containerNo'] or '',    # I: 柜号
            record['description'] or '',    # J: 异常描述
            record['responsible'] or '',    # K: 责任方
            record['fee'] or '',            # L: 异常费用
            record['progress'] or '',       # M: 处理进度
            record['status'] or '',         # N: 闭环情况
            record['agent'] or '',          # O: 代理
            ''                              # P: (empty, matches template)
        ]

        for col, val in enumerate(values, 1):
            cell = ws_abnormal.cell(row=row_num, column=col, value=val)
            style = abnormal_styles.get(col, {})
            if 'font' in style:
                cell.font = style['font']
            if 'alignment' in style:
                cell.alignment = style['alignment']
            if 'border' in style:
                cell.border = style['border']
            # Column B (发生/发现日期): use explicit date format
            if col == 2:
                cell.number_format = 'YYYY-MM-DD'
            elif 'number_format' in style:
                cell.number_format = style['number_format']

        # Auto-adjust row height based on the longest text column
        desc_lines = (record['description'] or '').count('\n') + 1
        prog_lines = (record['progress'] or '').count('\n') + 1
        max_lines = max(desc_lines, prog_lines, 1)
        row_height = max(15, max_lines * 15.6)
        ws_abnormal.row_dimensions[row_num].height = row_height

    # ============================================================
    # Merge column H (提运单号) cells for consecutive rows with the same value
    # Only merge 提单号 column; all other columns stay separate
    # ============================================================
    if len(abnormal_rows) >= 2:
        merge_start = 2  # first data row
        prev_bol = abnormal_rows[0]['bolNo'] or ''
        for i in range(1, len(abnormal_rows)):
            cur_bol = abnormal_rows[i]['bolNo'] or ''
            if cur_bol and prev_bol and cur_bol == prev_bol:
                # same 提单号 continues — extend merge range
                if i == len(abnormal_rows) - 1:
                    # last row, close merge
                    if merge_start < 2 + i:
                        ws_abnormal.merge_cells(
                            start_row=merge_start, end_row=2 + i,
                            start_column=8, end_column=8  # column H only
                        )
            else:
                # 提单号 changed — close previous merge if 2+ rows
                if merge_start < 2 + i - 1:
                    ws_abnormal.merge_cells(
                        start_row=merge_start, end_row=2 + i - 1,
                        start_column=8, end_column=8
                    )
                merge_start = 2 + i
                prev_bol = cur_bol

    # ============================================================
    # Merge column J (异常情况描述) for consecutive grouped records
    # Rule: first row has description + subsequent rows are empty +
    #        share same date + company + category
    # ============================================================
    if len(abnormal_rows) >= 2:
        i = 0
        while i < len(abnormal_rows):
            desc = (abnormal_rows[i]['description'] or '').strip()
            if not desc:
                i += 1
                continue

            # Found a row with description — look ahead for empty-description siblings
            group_start = i
            j = i + 1
            while j < len(abnormal_rows):
                cur_desc = (abnormal_rows[j]['description'] or '').strip()
                if (cur_desc == '' and
                    abnormal_rows[j]['date'] == abnormal_rows[group_start]['date'] and
                    (abnormal_rows[j]['category'] or '') == (abnormal_rows[group_start]['category'] or '')):
                    j += 1
                else:
                    break

            # Merge J column (col 10) if 2+ rows in group
            if j - group_start >= 2:
                ws_abnormal.merge_cells(
                    start_row=2 + group_start, end_row=2 + j - 1,
                    start_column=10, end_column=10
                )

            i = j  # Skip past the group

    # Remove conditional formatting from column H (提单号) — only keep on G (报关单号)
    cf_to_keep = [cf for cf in ws_abnormal.conditional_formatting
                  if not str(cf.sqref).replace('$', '').startswith('H')]
    ws_abnormal.conditional_formatting = cf_to_keep

    # Remove auto-filters from both sheets (exported file should show all data)
    # Clear ref AND filterColumn to ensure no filtering state persists
    ws_daily.auto_filter.ref = None
    ws_daily.auto_filter.filterColumn = ()
    ws_abnormal.auto_filter.ref = None
    ws_abnormal.auto_filter.filterColumn = ()

    # Also scrub autoFilter XML element if it lingers
    # (openpyxl < 3.1 uses _ws internal attribute; newer versions dont need this)
    try:
        for ws in (ws_daily, ws_abnormal):
            if not hasattr(ws, "_ws"):
                continue
            ns = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'
            for af in ws._ws.findall(f'{{{ns}}}autoFilter'):
                ws._ws.remove(af)
    except Exception:
        pass  # XML scrub is optional; ref=None + filterColumn=() already done above

    # ============================================================
    # Remove 2025 sheets and Sheet3 (only keep 2026 sheets)
    # ============================================================
    for sheet_name in ['2025 日报关情况', '2025年异常跟踪表', 'Sheet3']:
        if sheet_name in wb.sheetnames:
            del wb[sheet_name]

    # ============================================================
    # Save to BytesIO and return as download
    # ============================================================
    output = BytesIO()
    wb.save(output)
    output.seek(0)

    # Generate filename based on requested date (default: yesterday)
    yesterday = datetime.now() - timedelta(days=1)
    target_date = request.args.get('date', yesterday.strftime('%Y-%m-%d'))
    try:
        dt = datetime.strptime(target_date, '%Y-%m-%d')
        filename = f'华东口岸申报日报关情况 {dt.month}.{dt.day}.xlsx'
    except ValueError:
        filename = '华东口岸申报日报关情况.xlsx'

    response = app.response_class(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={
            'Content-Disposition': f"attachment; filename*=UTF-8''{quote(filename)}"
        }
    )
    response.headers['Cache-Control'] = 'no-store'
    return response


@app.route('/api/export-text-report', methods=['GET'])
def export_text_report():
    """Export daily text report as DOCX file.

    Sections:
    1. Daily summary line (declared/released/review counts)
    2. 【查验情况】— inspection records for the target date
    3. 【异常情况】— all unresolved (未闭环) abnormal records
    """
    from io import BytesIO
    from urllib.parse import quote

    db = get_db()

    # Parse target date (default: yesterday)
    yesterday = datetime.now() - timedelta(days=1)
    target_date = request.args.get('date', yesterday.strftime('%Y-%m-%d'))
    try:
        dt = datetime.strptime(target_date, '%Y-%m-%d')
    except ValueError:
        return jsonify({'error': 'Invalid date format, expected YYYY-MM-DD'}), 400

    month_day = f"{dt.month}月{dt.day}日"

    # 1. Daily record
    daily = db.execute(
        "SELECT * FROM daily_records WHERE date = ? AND year = '2026'",
        (target_date,)
    ).fetchone()

    total_declared = daily['totalDeclared'] if daily else 0
    total_released = daily['totalReleased'] if daily else 0
    review_completed = daily['reviewCompleted'] if daily else 0

    # If reviewCompleted is a count from unclearedReason text, use it;
    # otherwise fall back to: totalDeclared - totalReleased
    if review_completed <= 0 and daily and daily['unclearedReason']:
        import re
        m = re.search(r'(\d+)票审结未放行', daily['unclearedReason'] or '')
        if m:
            review_completed = int(m.group(1))

    # 2. Inspection records: today's 查验 + all ongoing 查验 (未闭环)
    inspection_records = db.execute("""
        SELECT * FROM abnormal_records
        WHERE year = '2026' AND deleted = 0
          AND category LIKE '%查验%'
          AND (date = ? OR status IN ('未闭环', '未关闭'))
        ORDER BY date, seq
    """, (target_date,)).fetchall()

    # 3. All unresolved (未闭环) records, excluding those with 查验 category
    abnormal_records = db.execute("""
        SELECT * FROM abnormal_records
        WHERE year = '2026' AND deleted = 0
          AND status IN ('未闭环', '未关闭')
          AND category NOT LIKE '%查验%'
        ORDER BY date, seq
    """).fetchall()

    # ============================================================
    # Build DOCX document
    # ============================================================
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        return jsonify({'error': 'python-docx 未安装，请运行: pip3 install python-docx'}), 500

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # === Title line: "6月5日" ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(month_day)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # === Summary line ===
    summary = doc.add_paragraph()
    summary_text = f"华东地区出口通关申报放行单量合共{total_released}单。审结{review_completed}票"
    run = summary.add_run(summary_text)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    note = doc.add_paragraph()
    run = note.add_run("当天走船计划除如下查验日报提及的单外已全部申报放行")
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # === 【查验情况】 ===
    doc.add_paragraph()  # blank line
    section_title = doc.add_paragraph()
    run = section_title.add_run("【查验情况】")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    if inspection_records:
        for idx, rec in enumerate(inspection_records, 1):
            item_para = doc.add_paragraph()
            item_para.paragraph_format.space_before = Pt(6)
            item_para.paragraph_format.space_after = Pt(2)

            lines = []
            lines.append(f"{idx}.经营单位：{rec['company'] or ''}")
            bol_part = f"/{rec['bolNo']}" if rec['bolNo'] else ""
            lines.append(f"报关单号/提单号：{rec['customsNo'] or ''}{bol_part}")
            lines.append(f"查验情况：{rec['description'] or ''}")
            lines.append(f"查验处理进度：{rec['progress'] or ''}")

            for line in lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(18)
                run = p.add_run(line)
                run.font.size = Pt(11)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            # Add blank line between items
            if idx < len(inspection_records):
                doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        run = p.add_run("无")
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # === 【异常情况】 ===
    doc.add_paragraph()
    section_title2 = doc.add_paragraph()
    run = section_title2.add_run("【异常情况】")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    if abnormal_records:
        for idx, rec in enumerate(abnormal_records, 1):
            lines = []
            lines.append(f"{idx}.业务类型：{rec['responsible'] or ''}")
            lines.append(f"经营单位：{rec['company'] or ''}")
            bol_part = f"/{rec['bolNo']}" if rec['bolNo'] else ""
            lines.append(f"报关单号/提单号：{rec['customsNo'] or ''}{bol_part}")
            lines.append(f"原因：{rec['description'] or ''}")
            lines.append(f"进度：{rec['progress'] or ''}")

            for line in lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(18)
                run = p.add_run(line)
                run.font.size = Pt(11)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            # Add blank line between items
            if idx < len(abnormal_records):
                doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        run = p.add_run("无")
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ============================================================
    # Save and return
    # ============================================================
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    filename = f'华东口岸申报日报关情况(文字版) {dt.month}.{dt.day}.docx'

    response = app.response_class(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Content-Disposition': f"attachment; filename*=UTF-8''{quote(filename)}"
        }
    )
    response.headers['Cache-Control'] = 'no-store'
    return response


# ---------------- Dimension Item Editing ----------------

@app.route('/api/daily-records/<record_id>/release-item', methods=['POST'])
def release_dimension_item(record_id):
    """Mark a dimension item as released — remove from unclearedReason, increment released count."""
    data = request.get_json(silent=True)
    if not data or 'item' not in data or 'dimension' not in data:
        return jsonify({'error': 'item and dimension required'}), 400

    db = get_db()
    record = db.execute("SELECT * FROM daily_records WHERE id = ?", (record_id,)).fetchone()
    if not record:
        return jsonify({'error': 'Daily record not found'}), 404

    old_reason = record['unclearedReason'] or ''
    item_text = str(data['item']).strip()
    dim_name = str(data['dimension']).strip()

    # Remove the specific item from the unclearedReason text
    new_reason = remove_item_from_reason(old_reason, dim_name, item_text)

    now = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
    new_released = (record['totalReleased'] or 0) + 1

    db.execute("""
        UPDATE daily_records SET
            unclearedReason = ?, totalReleased = ?,
            updated_at = ?, source = 'manual'
        WHERE id = ?
    """, (new_reason, new_released, now, record_id))
    db.commit()

    log_operation(
        action='release_item', record_type='daily', record_id=record_id,
        record_summary=f'放行: {dim_name} - {item_text}',
        changes={'dimension': dim_name, 'item': item_text},
        operator=get_operator()
    )

    return jsonify({'success': True, 'new_uncleared': new_reason, 'totalReleased': new_released})


@app.route('/api/daily-records/<record_id>/transfer-item', methods=['POST'])
def transfer_dimension_item(record_id):
    """Transfer a dimension item to abnormal records table."""
    data = request.get_json(silent=True)
    if not data or 'item' not in data or 'dimension' not in data:
        return jsonify({'error': 'item and dimension required'}), 400

    db = get_db()
    record = db.execute("SELECT * FROM daily_records WHERE id = ?", (record_id,)).fetchone()
    if not record:
        return jsonify({'error': 'Daily record not found'}), 404

    old_reason = record['unclearedReason'] or ''
    item_text = str(data['item']).strip()
    dim_name = str(data['dimension']).strip()

    # Remove from unclearedReason
    new_reason = remove_item_from_reason(old_reason, dim_name, item_text)

    now = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
    db.execute("""
        UPDATE daily_records SET unclearedReason = ?, updated_at = ? WHERE id = ?
    """, (new_reason, now, record_id))

    # Create abnormal record
    rid = str(uuid.uuid4())
    year = str(record['year'])
    max_seq = db.execute(
        "SELECT COALESCE(MAX(seq), 0) as mx FROM abnormal_records WHERE year = ?", (year,)
    ).fetchone()['mx']
    seq = max_seq + 1

    # Determine category from dimension
    cat_map = {'查验': '查验', '空运': '其他', '驳船': '其他', '大船': '其他', '公路': '其他'}
    category = cat_map.get(dim_name, '其他')

    # Parse item text: try to extract ID and reason
    parts = item_text.split(None, 1)
    item_id = parts[0] if parts else item_text
    item_reason = parts[1] if len(parts) > 1 else item_text

    db.execute("""
        INSERT INTO abnormal_records
        (id, seq, date, category, bizUnit, company, importExport,
         customsNo, bolNo, containerNo, description, responsible,
         fee, progress, status, agent, follow_up_notes, year,
         created_at, updated_at, source, deleted)
        VALUES (?, ?, ?, ?, '', '', '', '', ?, '', ?, '', '', '', '未闭环', '', '[]', ?, ?, ?, 'transfer', 0)
    """, (
        rid, seq, record['date'], category,
        item_id, item_reason,
        year, now, now
    ))
    db.commit()

    log_operation(
        action='transfer_item', record_type='daily', record_id=record_id,
        record_summary=f'转异常: {dim_name} - {item_text} → 异常记录 #{seq}',
        changes={'dimension': dim_name, 'item': item_text, 'new_abnormal_id': rid},
        operator=get_operator()
    )

    log_operation(
        action='create', record_type='abnormal', record_id=rid,
        record_summary=f'由未放行转入: {dim_name} {item_text}',
        changes={'source': 'dimension_transfer', 'dimension': dim_name},
        operator=get_operator()
    )

    return jsonify({
        'success': True,
        'new_uncleared': new_reason,
        'abnormal_id': rid,
        'abnormal_seq': seq
    })


@app.route('/api/daily-records/ensure', methods=['POST'])
def ensure_daily_record():
    """Ensure a daily record exists for the given date. Create one if not.
    Returns the record ID so callers can then add items to it."""
    data = request.get_json(silent=True)
    if not data or 'date' not in data:
        return jsonify({'error': 'date is required'}), 400

    db = get_db()
    rec_date = str(data['date']).strip()
    year = rec_date[:4]

    existing = db.execute(
        "SELECT id FROM daily_records WHERE date = ? AND year = ?",
        (rec_date, year)
    ).fetchone()

    if existing:
        return jsonify({'success': True, 'id': existing['id'], 'action': 'found'})

    # Create a stub daily record
    now = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
    rid = str(uuid.uuid4())
    dt = datetime.strptime(rec_date, '%Y-%m-%d')
    db.execute("""
        INSERT INTO daily_records
        (id, date, year, month, day, totalDeclared, totalReleased,
         reviewCompleted, unclearedReason, created_at, updated_at, source)
        VALUES (?, ?, ?, ?, ?, 0, 0, 0, '', ?, ?, 'manual')
    """, (rid, rec_date, year, dt.month, dt.day, now, now))
    db.commit()

    log_operation(
        action='ensure_daily', record_type='daily', record_id=rid,
        record_summary=f'自动创建日报: {rec_date}',
        changes={'date': rec_date},
        operator=get_operator()
    )

    return jsonify({'success': True, 'id': rid, 'action': 'created'})


@app.route('/api/daily-records/<record_id>/add-item', methods=['POST'])
def add_dimension_item(record_id):
    """Add a new item to a dimension section in the unclearedReason."""
    data = request.get_json(silent=True)
    if not data or 'dimension' not in data or 'item_id' not in data:
        return jsonify({'error': 'dimension and item_id required'}), 400

    db = get_db()
    record = db.execute("SELECT * FROM daily_records WHERE id = ?", (record_id,)).fetchone()
    if not record:
        return jsonify({'error': 'Daily record not found'}), 404

    old_reason = record['unclearedReason'] or ''
    dim_name = str(data['dimension']).strip()
    item_id = str(data['item_id']).strip()
    item_reason = str(data.get('reason', '')).strip()

    # Duplicate check against ALL daily records
    all_rows = db.execute(
        "SELECT id, date, unclearedReason FROM daily_records"
        " WHERE unclearedReason IS NOT NULL AND unclearedReason != '' AND unclearedReason != '无'"
    ).fetchall()
    for row in all_rows:
        if item_id in _extract_item_ids(row['unclearedReason'] or ''):
            return jsonify({
                'error': f'提单号重复，「{item_id}」已存在于 {row["date"]}',
                'duplicate': item_id
            }), 409

    new_reason = add_item_to_reason(old_reason, dim_name, item_id, item_reason)

    now = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
    new_review_completed = (record['reviewCompleted'] or 0) + 1

    db.execute("""
        UPDATE daily_records SET
            unclearedReason = ?, reviewCompleted = ?,
            updated_at = ?, source = 'manual'
        WHERE id = ?
    """, (new_reason, new_review_completed, now, record_id))
    db.commit()

    log_operation(
        action='add_item', record_type='daily', record_id=record_id,
        record_summary=f'新增未放行: {dim_name} - {item_id} {item_reason}',
        changes={'dimension': dim_name, 'item_id': item_id, 'reason': item_reason},
        operator=get_operator()
    )

    return jsonify({'success': True, 'new_uncleared': new_reason, 'reviewCompleted': new_review_completed})


def _parse_uncleared_sections(reason_str):
    """Parse unclearedReason text into {dim_name: [item_lines]} dict."""
    if not reason_str or reason_str == '无':
        return {}
    dim_keys = ['空运', '驳船', '大船', '公路', '查验']
    sections = {k: [] for k in dim_keys}
    current_dim = None
    for line in reason_str.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        is_header = False
        for dk in dim_keys:
            # Match "空运：", "空运:", or just "空运" (bare name, like frontend supports)
            if stripped.startswith(dk + '：') or stripped.startswith(dk + ':') or stripped == dk or stripped == dk:
                current_dim = dk
                is_header = True
                break
        if is_header:
            continue
        if re.match(r'^\d+票', stripped) or '审结未放行' in stripped:
            continue
        if current_dim:
            sections[current_dim].append(stripped)
    return {k: v for k, v in sections.items() if v}


def _rebuild_uncleared_text(sections):
    """Rebuild unclearedReason text from {dim_name: [item_lines]} dict."""
    total = sum(len(v) for v in sections.values())
    if total == 0:
        return '无'
    lines = [f"{total}票审结未放行 "]
    for dim_name in ['空运', '驳船', '大船', '公路', '查验']:
        items = sections.get(dim_name, [])
        if items:
            lines.append(f"{dim_name}：")
            lines.extend(items)
    return '\n'.join(lines)


def _extract_item_ids(reason_str):
    """Extract all item IDs (first token of each line) from unclearedReason. Returns a set."""
    sections = _parse_uncleared_sections(reason_str)
    ids = set()
    for items in sections.values():
        for item_line in items:
            parts = item_line.split(None, 1)
            if parts:
                ids.add(parts[0])
    return ids


def _merge_uncleared_reason(existing_str, new_str):
    """Merge new uncleared items into existing text by appending to each dimension section.
    Preserves original text format — does NOT rebuild, only appends truly new items."""
    if not new_str or new_str == '无':
        return existing_str or '无'
    if not existing_str or existing_str == '无':
        return new_str

    existing_sections = _parse_uncleared_sections(existing_str)
    new_sections = _parse_uncleared_sections(new_str)

    # Collect truly new items per dimension
    to_append = {}
    for dim_name in ['空运', '驳船', '大船', '公路', '查验']:
        existing_items = existing_sections.get(dim_name, [])
        new_items = new_sections.get(dim_name, [])
        existing_ids = {item.split(None, 1)[0] for item in existing_items}
        truly_new = [item for item in new_items if item.split(None, 1)[0] not in existing_ids]
        if truly_new:
            to_append[dim_name] = truly_new

    if not to_append:
        return existing_str  # No new items

    # Append new items to existing text at the correct dimension sections
    lines = existing_str.split('\n')
    result = []
    i = 0
    dim_keys = ['空运', '驳船', '大船', '公路', '查验']
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        # Check if this is a dimension header
        current_dim = None
        for dk in dim_keys:
            if stripped.startswith(dk + '：') or stripped.startswith(dk + ':') or stripped == dk:
                current_dim = dk
                break
        if current_dim and current_dim in to_append:
            result.append(line)  # Keep the header
            i += 1
            # Append existing items until next header or end
            while i < len(lines):
                next_stripped = lines[i].strip()
                is_next_header = False
                for dk in dim_keys:
                    if next_stripped.startswith(dk + '：') or next_stripped.startswith(dk + ':'):
                        is_next_header = True
                        break
                if is_next_header:
                    break
                if next_stripped:  # Non-empty item line
                    result.append(lines[i])
                i += 1
            # Append truly new items after existing items
            for new_item in to_append[current_dim]:
                result.append(new_item)
            del to_append[current_dim]  # Mark as done
        else:
            result.append(line)
            i += 1

    # If there are leftover dimensions not found in existing text, append at end
    for dim_name in ['空运', '驳船', '大船', '公路', '查验']:
        if dim_name in to_append:
            result.append(f"{dim_name}：")
            for new_item in to_append[dim_name]:
                result.append(new_item)

    return '\n'.join(result)


def remove_item_from_reason(reason_str, dimension, item_text):
    """Remove a specific item from an unclearedReason string for a given dimension.
    Also cleans up empty dimension headers."""
    if not reason_str or reason_str == '无':
        return '无'

    lines = reason_str.split('\n')
    result = []
    dim_header_indices = []  # Track (index, dim_name) for result lines
    pending_clear = set()    # Dimensions whose items have been fully cleared

    for line in lines:
        stripped = line.strip()
        is_header = False
        for dim_name in ['空运', '驳船', '大船', '公路', '查验']:
            if stripped.startswith(dim_name + '：') or stripped.startswith(dim_name + ':') or stripped == dim_name:
                is_header = True
                dim_header_indices.append((len(result), dim_name))
                break

        if is_header:
            result.append(line)
            continue

        # Check if this line contains the item to remove
        if item_text and (item_text in stripped or (item_text.split() and stripped.startswith(item_text.split()[0]))):
            continue  # Skip this item

        result.append(line)

    # Clean up empty dimension sections (header followed by no items or only whitespace until next header)
    cleaned = []
    i = 0
    while i < len(result):
        line = result[i]
        stripped = line.strip()
        # Check if this is a dimension header
        is_header = False
        for dim_name in ['空运', '驳船', '大船', '公路', '查验']:
            if stripped.startswith(dim_name + '：') or stripped.startswith(dim_name + ':') or stripped == dim_name:
                is_header = True
                break
        if is_header:
            # Look ahead: if next line is empty or another header, skip this header
            next_idx = i + 1
            has_content = False
            while next_idx < len(result):
                next_stripped = result[next_idx].strip()
                # Check if next line is another header
                next_is_header = False
                for dim_name in ['空运', '驳船', '大船', '公路', '查验']:
                    if next_stripped.startswith(dim_name + '：') or next_stripped.startswith(dim_name + ':'):
                        next_is_header = True
                        break
                if next_is_header:
                    break
                if next_stripped:
                    has_content = True
                    break
                next_idx += 1
            if not has_content:
                i += 1
                continue  # Skip empty dimension header
        cleaned.append(line)
        i += 1

    final = '\n'.join(cleaned).strip()
    return final if final else '无'


def add_item_to_reason(reason_str, dimension, item_id, reason):
    """Add an item to a specific dimension section in unclearedReason text.
    Creates the dimension section if it doesn't exist, updates the summary count."""
    import re
    dim_keys = ['空运', '驳船', '大船', '公路', '查验']
    new_item_line = f"{item_id} {reason}" if reason else item_id

    # Empty / no content — build from scratch
    if not reason_str or reason_str == '无':
        return f"1票审结未放行 \n{dimension}：\n{new_item_line}"

    lines = reason_str.split('\n')

    # Find target dimension header
    dim_idx = -1
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith(dimension + '：') or stripped.startswith(dimension + ':') or stripped == dimension:
            dim_idx = i
            break

    if dim_idx >= 0:
        # Find the next dimension header after dim_idx as the section boundary
        insert_idx = len(lines)
        for i in range(dim_idx + 1, len(lines)):
            stripped = lines[i].strip()
            for dk in dim_keys:
                if stripped.startswith(dk + '：') or stripped.startswith(dk + ':') or stripped == dk:
                    insert_idx = i
                    break
            if insert_idx < len(lines):
                break

        # Insert after the last non-empty line inside this section
        actual_insert = insert_idx
        for j in range(insert_idx - 1, dim_idx, -1):
            if lines[j].strip():
                actual_insert = j + 1
                break
        else:
            actual_insert = dim_idx + 1  # Empty section, insert right after header

        lines.insert(actual_insert, new_item_line)
    else:
        # Dimension section doesn't exist — append at end
        lines.append(f"{dimension}：")
        lines.append(new_item_line)

    # Increment the summary count line (e.g. "23票审结未放行" → "24票")
    for i, line in enumerate(lines):
        m = re.match(r'^(\d+)票审结未放行', line.strip())
        if m:
            count = int(m.group(1)) + 1
            lines[i] = line.replace(m.group(0), f"{count}票审结未放行", 1)
            break

    return '\n'.join(lines).strip()


# ---------------- Startup ----------------

# Initialize database and run migration on import
init_db()
migrate_from_json()

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8888))
    print(f"🚀 美的项目客服看板 v2 启动在端口 {port}")
    print(f"   数据库: {DB_FILE}")
    app.run(host='0.0.0.0', port=port, debug=False)
